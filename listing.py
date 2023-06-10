import os
import sys
import time
import PyQt6

import docx
from docx.shared import Pt
from docx.shared import Length
from docx.shared import RGBColor


from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtCore import QEvent, QObject
from PyQt6.QtGui import QStandardItemModel, QStandardItem
from PyQt6.QtWidgets import QFileDialog, QComboBox, QCheckBox

class CheckableComboBox(QtWidgets.QComboBox):
    popup = QtCore.pyqtSignal()
    states = []

    def __init__(self, parent=None):
        super(CheckableComboBox, self).__init__(parent)
        self.view().pressed.connect(self.handleItemPressed)
        self._changed = False

    def addItem(self, item):
        super().addItem(item)
        item = self.model().item(self.count()-1,0)
        item.setFlags(QtCore.Qt.ItemFlag.ItemIsUserCheckable | QtCore.Qt.ItemFlag.ItemIsEnabled)
        item.setCheckState(QtCore.Qt.CheckState.Unchecked)
        self._changed = True

    def addBaseItem(self, item):
        super().addItem(item)

    def itemChecked(self, index):
        item = self.model().item(index,0)
        return item.checkState() == QtCore.Qt.CheckState.Checked

    def setItemChecked(self, index, checked=True):
        item = self.model().item(index, 0)
        if checked:
            item.setCheckState(QtCore.Qt.CheckState.Checked)
        else:
            item.setCheckState(QtCore.Qt.CheckState.Unchecked)

    def showPopup(self) -> None:
        self.popup.emit()
        return super().showPopup()
    
    def hidePopup(self) -> None:
        self.states = []
        for i in range(self.count()):
            self.states.append(self.model().item(i,0).checkState() == QtCore.Qt.CheckState.Checked)
        print(self.states)
        return super().hidePopup()
    
    def handleItemPressed(self, index):
        item = self.model().itemFromIndex(index)
        if item.checkState() == QtCore.Qt.CheckState.Checked:
            item.setCheckState(QtCore.Qt.CheckState.Unchecked)
        else:
            item.setCheckState(QtCore.Qt.CheckState.Checked)



class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(600, 250)
        MainWindow.setStyleSheet("")
        MainWindow.setIconSize(QtCore.QSize(300, 400))
        self.centralwidget = QtWidgets.QWidget(parent=MainWindow)
        self.centralwidget.setObjectName("centralwidget")

        self.InputLine = QtWidgets.QLineEdit(parent=self.centralwidget)
        self.InputLine.setGeometry(QtCore.QRect(20, 80, 381, 28))
        self.InputLine.setObjectName("InputLine")
        self.InputLine.textChanged.connect(self.InlineChanged)

        self.InputBtn = QtWidgets.QPushButton(parent=self.centralwidget)
        self.InputBtn.setGeometry(QtCore.QRect(420, 80, 93, 28))
        font = QtGui.QFont()
        font.setFamily("Montserrat SemiBold")
        self.InputBtn.setFont(font)
        self.InputBtn.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.PointingHandCursor))
        self.InputBtn.setObjectName("InputBtn")
        self.InputBtn.clicked.connect(self.InputBtnClick)

        self.OutputLine = QtWidgets.QLineEdit(parent=self.centralwidget)
        self.OutputLine.setGeometry(QtCore.QRect(20, 130, 381, 28))
        self.OutputLine.setObjectName("OutputLine")

        self.OutputBtn = QtWidgets.QPushButton(parent=self.centralwidget)
        self.OutputBtn.setGeometry(QtCore.QRect(420, 130, 93, 28))
        font = QtGui.QFont()
        font.setFamily("Montserrat SemiBold")
        self.OutputBtn.setFont(font)
        self.OutputBtn.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.PointingHandCursor))
        self.OutputBtn.setObjectName("OutputBtn")
        self.OutputBtn.clicked.connect(self.OutputBtnClick)

        self.label = QtWidgets.QLabel(parent=self.centralwidget)
        self.label.setGeometry(QtCore.QRect(20, 60, 191, 16))
        font = QtGui.QFont()
        font.setFamily("Montserrat Medium")
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(parent=self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(20, 110, 211, 16))
        font = QtGui.QFont()
        font.setFamily("Montserrat Medium")
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(parent=self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(0, 0, 600, 50))
        font = QtGui.QFont()
        font.setFamily("Montserrat SemiBold")
        font.setPointSize(20)
        self.label_3.setFont(font)
        self.label_3.setStyleSheet("")
        self.label_3.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(parent=self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(20, 170, 151, 16))
        font = QtGui.QFont()
        font.setFamily("Montserrat Medium")
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.FormBtn = QtWidgets.QPushButton(parent=self.centralwidget)
        self.FormBtn.setGeometry(QtCore.QRect(250, 170, 301, 41))
        font = QtGui.QFont()
        font.setFamily("Montserrat ExtraBold")
        font.setPointSize(15)
        self.FormBtn.setFont(font)
        self.FormBtn.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.PointingHandCursor))
        self.FormBtn.setStyleSheet("background-color: rgb(170, 0, 0);\n"
            "color: rgb(255, 255, 255);")
        self.FormBtn.setObjectName("FormBtn")
        self.FormBtn.clicked.connect(self.FormClick)

        self.comboBox = CheckableComboBox(self.centralwidget)
        self.comboBox.setDuplicatesEnabled(False)
        self.comboBox.addBaseItem("")
        self.comboBox.setGeometry(QtCore.QRect(20, 190, 181, 22))
        self.comboBox.setObjectName("comboBox")
        self.comboBox.popup.connect(self.ComboBoxActivate)
        self.comboBox.view().pressed.connect(self.comboBox.handleItemPressed)
        
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Listing"))
        self.InputBtn.setText(_translate("MainWindow", "Open"))
        self.OutputBtn.setText(_translate("MainWindow", "Open"))
        self.label.setText(_translate("MainWindow", "Input directory"))
        self.label_2.setText(_translate("MainWindow", "Output directory"))
        self.label_3.setText(_translate("MainWindow", "Form the listing of program"))
        self.label_4.setText(_translate("MainWindow", "Format of the files"))
        self.FormBtn.setText(_translate("MainWindow", "Form"))

    def InputBtnClick(self):
        path = QFileDialog.getExistingDirectoryUrl(
           None, "Input Directory").toString().replace("file:///", "")
        
        self.InputLine.setText(path)

    def OutputBtnClick(self):
        path = QFileDialog.getExistingDirectoryUrl(
           None, "Input Directory").toString().replace("file:///", "")
        
        self.OutputLine.setText(path)
    
    def ComboBoxActivate(self):
        self.comboBox.clear()
        try:
            file_types = []
            for file_name in os.listdir(self.InputLine.text()):
                if file_name.find(".") != -1:
                    file_types.append(file_name[file_name.index(".") : len(file_name)])
            
            file_types = list(set(file_types))

            counter = 0

            for file_type in file_types:
                self.comboBox.addItem(file_type)
                self.comboBox.setItemChecked(counter, False)
                counter += 1
        except FileNotFoundError:
            self.comboBox.addBaseItem("")
    
    def InlineChanged(self):
        self.comboBox.clear()

    def FormClick(self):
        if self.InputLine.text() == "":
            self.InputLine.setText("Error: no input path")
            return False
        elif self.OutputLine.text() == "":
            self.OutputLine.setText("Error: no output path")
            return False
        
        file_types = []
        for i in range(self.comboBox.count()):
            if self.comboBox.itemChecked(i):
                file_types.append(self.comboBox.model().item(i,0).text())
        
        if file_types == []:
            self.label_3.setText("Error: no chosen format")
            return False
        
        make_listing(self.InputLine.text(), self.OutputLine.text(), file_types)
        self.label_3.setText("Successful forming listing:)")



def make_listing(input_path, output_path, file_types):

    doc = docx.Document()

    style = doc.styles["Heading3"]
    style.font.name = "Times New Romans"
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    for file_type in file_types:
        for file_name in os.listdir(input_path):
            if file_name.find(".") != -1:
                if file_name[file_name.index(".") : len(file_name)] == file_type:
                        doc.add_heading("Код файла " + file_name, 3)
                        with open(input_path + "/" + file_name, "r") as file:
                            for line in file.readlines():
                                paragraph = doc.add_paragraph(line.replace("\n", ""))
                                paragraph.paragraph_format.line_spacing = Pt(18)

    doc.save(output_path + "/listing.docx")



if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    app.setWindowIcon(QtGui.QIcon('icon.ico'))

    MainWindow = QtWidgets.QMainWindow()
    MainWindow.setWindowIcon(QtGui.QIcon('icon.ico'))

    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)

    MainWindow.show()

    sys.exit(app.exec())


